# app/services/ocr_service.py
import io
import logging
import os
import tempfile
from pathlib import Path
from typing import Optional

# ── Third-party: python-docx ──────────────────────────────────────
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ── Third-party: Pillow ───────────────────────────────────────────
from PIL import Image, ImageEnhance, ImageFilter

# ── Internal: your app config ─────────────────────────────────────
from app.config import settings, GOOGLE_VISION_ENABLED



from docx.shared import Pt, Inches


logger = logging.getLogger(__name__)


# ─────────────────────────────────────────
#  IMAGE PREPROCESSING (improves accuracy)
# ─────────────────────────────────────────

def _preprocess_image(image: Image.Image) -> Image.Image:
    """
    Enhance image quality before sending to OCR.
    Improves accuracy on scanned/low-quality images.
    """
    # Convert to RGB if needed (handles RGBA PNGs, etc.)
    if image.mode not in ("RGB", "L"):
        image = image.convert("RGB")

    # Upscale if too small (Vision API struggles below 300px)
    min_dimension = 1000
    if min(image.size) < min_dimension:
        scale = min_dimension / min(image.size)
        new_size = (int(image.size[0] * scale), int(image.size[1] * scale))
        image = image.resize(new_size, Image.LANCZOS)

    # Enhance contrast for faded/scanned documents
    image = ImageEnhance.Contrast(image).enhance(1.5)

    # Slight sharpening
    image = image.filter(ImageFilter.SHARPEN)

    return image


def _image_to_bytes(image: Image.Image) -> bytes:
    """Convert PIL Image to PNG bytes for Vision API."""
    buffer = io.BytesIO()
    image.save(buffer, format="PNG")
    return buffer.getvalue()


# ─────────────────────────────────────────
#  PDF → LIST OF IMAGES
# ─────────────────────────────────────────

def _pdf_to_images(file_bytes: bytes) -> list[Image.Image]:
    """
    Convert each PDF page into a PIL Image using PyMuPDF.
    No poppler needed — works on Render, Railway, any platform.
    """
    import fitz  # PyMuPDF — already in your requirements.txt

    images = []
    pdf_document = fitz.open(stream=file_bytes, filetype="pdf")

    for page_num in range(len(pdf_document)):
        page = pdf_document[page_num]

        # 300 DPI equivalent — matrix scale of 300/72 ≈ 4.17
        mat = fitz.Matrix(300 / 72, 300 / 72)
        pix = page.get_pixmap(matrix=mat, alpha=False)

        # Convert PyMuPDF pixmap → PIL Image
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        images.append(img)
        logger.info(f"Rendered PDF page {page_num + 1}/{len(pdf_document)}")

    pdf_document.close()

    if not images:
        raise ValueError("PDF appears to have no renderable pages.")

    return images

# ─────────────────────────────────────────
#  GOOGLE VISION OCR
# ─────────────────────────────────────────

def _ocr_with_google_vision(image_bytes: bytes) -> str:
    """
    Send a single image to Google Cloud Vision API.
    Returns the extracted plain text.
    """
    from google.cloud import vision

    client = vision.ImageAnnotatorClient()
    image = vision.Image(content=image_bytes)

    # DOCUMENT_TEXT_DETECTION is better than TEXT_DETECTION
    # for multi-paragraph documents — preserves line breaks better
    response = client.document_text_detection(image=image)

    if response.error.message:
        raise RuntimeError(f"Google Vision API error: {response.error.message}")

    if response.full_text_annotation:
        return response.full_text_annotation.text

    return ""


# ─────────────────────────────────────────
#  TESSERACT FALLBACK
# ─────────────────────────────────────────

def _ocr_with_tesseract(image: Image.Image) -> str:
    """
    Fallback OCR using Tesseract (your existing engine).
    Used when Google Vision is disabled or fails.
    """
    import pytesseract
    try:
        import pytesseract
        return pytesseract.image_to_string(image, lang="eng")
    except Exception as e:
        logger.error(f"Tesseract OCR failed: {e}")
        return ""


# ─────────────────────────────────────────
#  OCR DISPATCHER
# ─────────────────────────────────────────

def _run_ocr_on_image(image: Image.Image) -> str:
    """
    Route to the best available OCR engine.
    Priority: Google Vision → Tesseract fallback
    """
    processed = _preprocess_image(image)
    image_bytes = _image_to_bytes(processed)

    if GOOGLE_VISION_ENABLED:
        try:
            logger.info("Using Google Vision API for OCR")
            return _ocr_with_google_vision(image_bytes)
        except Exception as e:
            logger.warning(f"Google Vision failed, falling back to Tesseract: {e}")

    logger.info("Using Tesseract for OCR")
    return _ocr_with_tesseract(processed)


# ─────────────────────────────────────────
#  .DOCX BUILDER
# ─────────────────────────────────────────

def _build_docx(pages_text: list[str], source_filename: str) -> bytes:
    """
    Convert extracted text pages into a formatted .docx file.
    Each PDF page gets its own section with a page break.
    """
    doc = Document()

    # Document title
    title = doc.add_heading(f"OCR Output — {source_filename}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for page_num, text in enumerate(pages_text, start=1):
        # Page label (for multi-page PDFs)
        if len(pages_text) > 1:
            page_heading = doc.add_heading(f"Page {page_num}", level=2)
            page_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Split into paragraphs by double newlines
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]

        if not paragraphs:
            # If no double-newlines found, treat single lines as paragraphs
            paragraphs = [line.strip() for line in text.split("\n") if line.strip()]

        for para_text in paragraphs:
            para = doc.add_paragraph(para_text)
            para.style.font.size = Pt(11)

        # Page break between PDF pages (skip after last page)
        if page_num < len(pages_text):
            doc.add_page_break()

    # Return as bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ─────────────────────────────────────────
#  MAIN PUBLIC FUNCTION
#  (called from documents.py router)
# ─────────────────────────────────────────

def perform_ocr_and_convert(
    file_bytes: bytes,
    filename: str,
    mime_type: str,
) -> tuple[bytes, str]:
    """
    Main OCR entry point. Accepts image or PDF bytes.

    Returns:
        tuple: (docx_bytes, output_filename)

    Raises:
        ValueError: If file type is unsupported
        RuntimeError: If OCR fails on all engines
    """
    pages_text: list[str] = []
    stem = Path(filename).stem

    # ── PDF input ──────────────────────────────────────
    if mime_type == "application/pdf" or filename.lower().endswith(".pdf"):
        logger.info(f"Processing PDF: {filename}")
        images = _pdf_to_images(file_bytes)

        if not images:
            raise ValueError("Could not extract pages from PDF.")

        for i, page_image in enumerate(images):
            logger.info(f"OCR on page {i+1}/{len(images)}")
            text = _run_ocr_on_image(page_image)
            pages_text.append(text)

    # ── Image input (PNG, JPG, TIFF, WEBP, BMP) ───────
    elif mime_type.startswith("image/") or filename.lower().endswith(
        (".png", ".jpg", ".jpeg", ".tiff", ".tif", ".webp", ".bmp")
    ):
        logger.info(f"Processing image: {filename}")
        image = Image.open(io.BytesIO(file_bytes))
        text = _run_ocr_on_image(image)
        pages_text.append(text)

    else:
        raise ValueError(f"Unsupported file type for OCR: {mime_type}")

    # Check we actually got something
    combined_text = " ".join(pages_text).strip()
    if not combined_text:
        raise RuntimeError(
            "OCR produced no text. The document may be blank, "
            "image-only without text, or too low quality."
        )

    # Build the .docx
    docx_bytes = _build_docx(pages_text, stem)
    output_filename = f"{stem}_ocr.docx"

    return docx_bytes, output_filename